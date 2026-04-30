"""
Universal File Converter
Run: python file_converter.py

Conversions supported:
  - CSV / XLSX  → PNG / JPG / JPEG
  - PDF         → PNG / JPG / JPEG  (one image per page)
  - Images      → PDF               (multiple images merged)
  - DOCX        → PDF
"""

import os
import sys
import threading
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from pathlib import Path

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt


# ═══════════════════════════════════════════════════════════
#  BUNDLED POPPLER PATH (for PyInstaller .exe)
# ═══════════════════════════════════════════════════════════

if getattr(sys, 'frozen', False):
    # running as .exe — use the bundled poppler inside the package
    BASE_DIR = sys._MEIPASS
else:
    # running as .py — look for poppler/ next to this script
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

POPPLER_PATH = os.path.join(BASE_DIR, "poppler")
if not os.path.isdir(POPPLER_PATH):
    POPPLER_PATH = None   # fall back to system PATH


# ═══════════════════════════════════════════════════════════
#  CONVERSION FUNCTIONS
# ═══════════════════════════════════════════════════════════

def spreadsheet_to_image(src, fmt, log):
    from PIL import Image, ImageOps
    import io

    ext = Path(src).suffix.lower()
    df  = pd.read_csv(src) if ext == ".csv" else pd.read_excel(src)
    rows, cols = df.shape
    log(f"Loaded {rows} rows x {cols} columns")

    col_width = max(1.4, min(2.5, 12 / max(cols, 1)))
    fig_w = max(8,  cols * col_width + 1)
    fig_h = max(2, (rows + 1) * 0.38)   # tighter height, no title padding

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    fig.subplots_adjust(left=0, right=1, top=1, bottom=0)  # zero margins
    ax.axis("off")

    row_fill = [["#F7F9FC" if i % 2 == 0 else "#FFFFFF"] * cols for i in range(rows)]
    tbl = ax.table(
        cellText    = df.astype(str).values.tolist(),
        colLabels   = df.columns.tolist(),
        cellLoc     = "center",
        loc         = "center",
        cellColours = row_fill,
        colColours  = ["#2E4057"] * cols,
    )
    tbl.auto_set_font_size(False)
    font_size = max(7, min(11, int(120 / max(cols, rows, 1))))
    tbl.set_fontsize(font_size)
    tbl.auto_set_column_width(col=list(range(cols)))

    # scale table to fill the axes exactly
    tbl.scale(1, 1.4)

    for (r, c), cell in tbl.get_celld().items():
        cell.set_edgecolor("#D0D7E3")
        cell.set_linewidth(0.6)
        cell.get_text().set_color("#FFFFFF" if r == 0 else "#2C2C2C")
        if r == 0:
            cell.get_text().set_fontweight("bold")

    # save to an in-memory buffer first
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, facecolor="white",
                bbox_inches="tight", pad_inches=0)
    plt.close(fig)

    # crop any remaining whitespace with Pillow
    buf.seek(0)
    img = Image.open(buf).convert("RGB")
    # invert to find non-white pixels and get bounding box
    inverted = ImageOps.invert(img)
    bbox = inverted.getbbox()          # (left, top, right, bottom)
    if bbox:
        padding = 6                    # small padding around table in pixels
        left   = max(0,          bbox[0] - padding)
        top    = max(0,          bbox[1] - padding)
        right  = min(img.width,  bbox[2] + padding)
        bottom = min(img.height, bbox[3] + padding)
        img = img.crop((left, top, right, bottom))

    save_fmt = "jpeg" if fmt in ("jpg", "jpeg") else "png"
    out = str(Path(src).with_suffix(f".{fmt}"))
    img.save(out, save_fmt.upper(), dpi=(150, 150))

    log(f"Saved: {out}")
    return [out]


def pdf_to_images(src, fmt, log):
    try:
        from pdf2image import convert_from_path
    except ImportError:
        raise RuntimeError("pdf2image not installed. Run: pip install pdf2image")

    log("Converting PDF pages to images...")
    kwargs = {"dpi": 150}
    if POPPLER_PATH:
        kwargs["poppler_path"] = POPPLER_PATH

    images = convert_from_path(src, **kwargs)
    log(f"Found {len(images)} page(s)")

    out_dir = Path(src).parent
    stem    = Path(src).stem
    pil_fmt = "JPEG" if fmt in ("jpg", "jpeg") else "PNG"
    ext     = "jpg"  if fmt in ("jpg", "jpeg") else "png"
    saved   = []
    for i, img in enumerate(images, 1):
        out = str(out_dir / f"{stem}_page{i}.{ext}")
        img.save(out, pil_fmt)
        log(f"Saved: {Path(out).name}")
        saved.append(out)
    return saved


def images_to_pdf(src_list, log):
    try:
        import img2pdf
    except ImportError:
        raise RuntimeError("img2pdf not installed. Run: pip install img2pdf")

    from PIL import Image

    log(f"Merging {len(src_list)} image(s) into PDF...")
    tmp_files = []
    converted = []

    for p in src_list:
        try:
            img = Image.open(p)
            img.load()
        except Exception as e:
            raise RuntimeError(f"Cannot open image: {Path(p).name}\n{e}")

        if img.mode in ("RGBA", "PA", "LA", "P"):
            img = img.convert("RGBA")
            tmp = p + "__tmp.png"
            img.save(tmp, "PNG")
        else:
            img = img.convert("RGB")
            tmp = p + "__tmp.jpg"
            img.save(tmp, "JPEG", quality=95)

        tmp_files.append(tmp)
        converted.append(tmp)
        log(f"Prepared: {Path(p).name}")

    if not converted:
        raise RuntimeError("No valid images to convert.")

    out = str(Path(src_list[0]).parent / "images_merged.pdf")
    pdf_bytes = img2pdf.convert(converted)

    if pdf_bytes is None:
        raise RuntimeError(
            "img2pdf returned no data. "
            "Try converting your images to JPG first and try again."
        )

    with open(out, "wb") as f:
        f.write(pdf_bytes)

    for t in tmp_files:
        try:
            os.remove(t)
        except Exception:
            pass

    log(f"Saved: {out}")
    return [out]


def _find_libreoffice():
    """Find LibreOffice executable across Windows, Mac and Linux."""
    import shutil, glob

    # check PATH first (works on Mac/Linux if properly installed)
    for cmd in ("libreoffice", "soffice"):
        path = shutil.which(cmd)
        if path:
            return path

    # common Windows install locations
    win_patterns = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice*\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice*\program\soffice.exe",
    ]
    for pattern in win_patterns:
        matches = glob.glob(pattern)
        if matches:
            return matches[0]

    # common Mac install location
    mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.exists(mac_path):
        return mac_path

    return None


def docx_to_pdf(src, log):
    log("Converting DOCX to PDF...")
    import subprocess

    out = str(Path(src).with_suffix(".pdf"))

    # ── Method 1: LibreOffice (best quality) ──────────────────
    lo = _find_libreoffice()
    if lo:
        log(f"Using LibreOffice...")
        result = subprocess.run(
            [lo, "--headless", "--convert-to", "pdf",
             "--outdir", str(Path(src).parent), src],
            capture_output=True, text=True
        )
        if result.returncode == 0 and os.path.exists(out):
            log(f"Saved: {out}")
            return [out]
        else:
            log(f"LibreOffice error: {result.stderr.strip() or result.stdout.strip()}")

    # ── Method 2: docx2pdf (Windows/Mac with Word installed) ──
    try:
        from docx2pdf import convert as d2p_convert
        log("Using docx2pdf (Microsoft Word)...")
        d2p_convert(src, out)
        if os.path.exists(out):
            log(f"Saved: {out}")
            return [out]
    except Exception:
        pass

    # ── Method 3: Pure Python fallback ────────────────────────
    try:
        log("Using built-in converter (install LibreOffice for best results)...")
        _docx_to_pdf_pure(src, out, log)
        log(f"Saved: {out}")
        return [out]
    except Exception as e:
        raise RuntimeError(
            f"Conversion failed: {e}\n\n"
            "Please install LibreOffice for best results:\n"
            "https://www.libreoffice.org/download/download-libreoffice"
        )


def _docx_to_pdf_pure(src, out, log):
    """Pure Python DOCX → PDF — faithful reproduction using python-docx + reportlab."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, pt
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

    PAGE_W, PAGE_H = A4
    L_MARGIN = R_MARGIN = 2.5 * cm
    USABLE_W = PAGE_W - L_MARGIN - R_MARGIN

    doc = Document(src)
    pdf = SimpleDocTemplate(
        out, pagesize=A4,
        leftMargin=L_MARGIN, rightMargin=R_MARGIN,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )

    base_styles = getSampleStyleSheet()

    def make_style(name, parent="Normal", **kw):
        return ParagraphStyle(name, parent=base_styles[parent], **kw)

    h1 = make_style("H1", "Heading1", fontSize=16, spaceAfter=8,  spaceBefore=12, textColor=colors.HexColor("#1a1a2e"))
    h2 = make_style("H2", "Heading2", fontSize=13, spaceAfter=6,  spaceBefore=10, textColor=colors.HexColor("#2e4057"))
    h3 = make_style("H3", "Heading3", fontSize=11, spaceAfter=4,  spaceBefore=8,  textColor=colors.HexColor("#374151"))
    normal      = make_style("Body",   fontSize=10, leading=14, spaceAfter=4)
    bold_normal = make_style("BodyB",  fontSize=10, leading=14, spaceAfter=4, fontName="Helvetica-Bold")
    bullet_st   = make_style("Bullet", fontSize=10, leading=14, spaceAfter=3, leftIndent=18, bulletIndent=6)

    align_map = {"LEFT": TA_LEFT, "CENTER": TA_CENTER, "RIGHT": TA_RIGHT, "JUSTIFY": TA_JUSTIFY}

    # ── helpers ────────────────────────────────────────────────────────────

    def escape(t):
        return t.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    def runs_to_html(para):
        """Convert paragraph runs to HTML-like string respecting bold/italic/underline/color."""
        parts = []
        for run in para.runs:
            txt = escape(run.text)
            if not txt:
                continue
            # detect bold: explicit run bold OR style bold
            is_bold = run.bold or (run.style and run.style.font.bold)
            is_italic = run.italic or (run.style and run.style.font.italic)
            is_under  = run.underline

            # font color
            try:
                rgb = run.font.color.rgb
                txt = f'<font color="#{rgb}">{txt}</font>'
            except Exception:
                pass

            if is_bold:    txt = f"<b>{txt}</b>"
            if is_italic:  txt = f"<i>{txt}</i>"
            if is_under:   txt = f"<u>{txt}</u>"
            parts.append(txt)
        return "".join(parts)

    def para_is_bold(p):
        """True if the whole paragraph is effectively bold."""
        if p.runs and all(r.bold for r in p.runs if r.text.strip()):
            return True
        try:
            return p.style.font.bold
        except Exception:
            return False

    def get_para_style(p):
        name = (p.style.name or "").strip()
        if name.startswith("Heading 1") or name == "Title": return h1
        if name.startswith("Heading 2"):                    return h2
        if name.startswith("Heading 3"):                    return h3
        if "List"  in name:                                 return bullet_st
        align_raw = str(p.alignment).upper() if p.alignment else "LEFT"
        align     = align_map.get(align_raw, TA_LEFT)
        if para_is_bold(p):
            return make_style(f"bb_{id(p)}", "Normal",
                              fontSize=10, leading=14, spaceAfter=4,
                              fontName="Helvetica-Bold", alignment=align)
        return make_style(f"np_{id(p)}", "Normal",
                          fontSize=10, leading=14, spaceAfter=4, alignment=align)

    def process_paragraph(p):
        html = runs_to_html(p)
        name = (p.style.name or "").strip()
        if not html.strip():
            return Spacer(1, 4)
        prefix = "• " if "List" in name else ""
        return Paragraph(prefix + html, get_para_style(p))

    # ── table processing ───────────────────────────────────────────────────

    def cell_runs_to_html(cell):
        parts = []
        for p in cell.paragraphs:
            parts.append(runs_to_html(p))
        return "<br/>".join(p for p in parts if p)

    def cell_is_bold(cell):
        for p in cell.paragraphs:
            if para_is_bold(p):
                return True
            for run in p.runs:
                if run.bold:
                    return True
        return False

    def get_col_widths(tbl):
        """Try to read real column widths from the XML; fall back to equal split."""
        try:
            tbl_w_el = tbl._tbl.find(qn("w:tblPr"), tbl._tbl.nsmap)
            grid = tbl._tbl.findall(qn("w:tblGrid") + "/" + qn("w:gridCol"),
                                    tbl._tbl.nsmap)
            if not grid:
                grid = tbl._tbl.find(qn("w:tblGrid"), tbl._tbl.nsmap)
                if grid is not None:
                    grid = grid.findall(qn("w:gridCol"))
            if grid:
                twips = [int(g.get(qn("w:w"), 0)) for g in grid]
                total = sum(twips) or 1
                return [USABLE_W * (w / total) for w in twips]
        except Exception:
            pass
        n = len(tbl.columns)
        return [USABLE_W / n] * n

    def process_table(tbl):
        col_widths = get_col_widths(tbl)
        n_cols     = len(col_widths)

        # build cell data as Paragraphs + collect span commands
        data       = []
        style_cmds = [
            ("GRID",         (0,0), (-1,-1), 0.5, colors.HexColor("#AAAAAA")),
            ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
            ("LEFTPADDING",  (0,0), (-1,-1), 5),
            ("RIGHTPADDING", (0,0), (-1,-1), 5),
            ("TOPPADDING",   (0,0), (-1,-1), 4),
            ("BOTTOMPADDING",(0,0), (-1,-1), 4),
            ("FONTSIZE",     (0,0), (-1,-1), 9),
        ]

        seen_spans = set()   # track cells that are part of a span

        for ri, row in enumerate(tbl.rows):
            row_data = []
            ci = 0
            for cell in row.cells:
                # detect merged cells via tc element identity
                tc = cell._tc
                cell_key = id(tc)

                if cell_key in seen_spans:
                    row_data.append("")   # placeholder for spanned cell
                    ci += 1
                    continue

                # check gridSpan (horizontal merge)
                grid_span = 1
                tc_pr = tc.find(qn("w:tcPr"))
                if tc_pr is not None:
                    gs_el = tc_pr.find(qn("w:gridSpan"))
                    if gs_el is not None:
                        grid_span = int(gs_el.get(qn("w:val"), 1))

                # check vMerge (vertical merge)
                v_merge = None
                if tc_pr is not None:
                    vm_el = tc_pr.find(qn("w:vMerge"))
                    if vm_el is not None:
                        v_merge = vm_el.get(qn("w:val"), "continue")

                # build cell content
                html = cell_runs_to_html(cell)
                is_bold = cell_is_bold(cell)

                cell_style = ParagraphStyle(
                    f"cell_{ri}_{ci}",
                    parent=base_styles["Normal"],
                    fontSize=9, leading=13,
                    fontName="Helvetica-Bold" if is_bold else "Helvetica",
                )

                # detect cell background color
                try:
                    shd = tc_pr.find(qn("w:shd")) if tc_pr is not None else None
                    if shd is not None:
                        fill = shd.get(qn("w:fill"), "")
                        if fill and fill != "auto" and len(fill) == 6:
                            bg = colors.HexColor(f"#{fill}")
                            style_cmds.append(("BACKGROUND", (ci, ri), (ci + grid_span - 1, ri), bg))
                            # white text on dark backgrounds
                            r2, g2, b2 = int(fill[0:2],16), int(fill[2:4],16), int(fill[4:6],16)
                            if (r2*299 + g2*587 + b2*114) / 1000 < 128:
                                cell_style = ParagraphStyle(
                                    f"cellW_{ri}_{ci}",
                                    parent=base_styles["Normal"],
                                    fontSize=9, leading=13,
                                    fontName="Helvetica-Bold" if is_bold else "Helvetica",
                                    textColor=colors.white,
                                )
                                style_cmds.append(("TEXTCOLOR", (ci, ri), (ci + grid_span - 1, ri), colors.white))
                except Exception:
                    pass

                para = Paragraph(html, cell_style) if html.strip() else ""
                row_data.append(para)

                # register horizontal span
                if grid_span > 1:
                    style_cmds.append(("SPAN", (ci, ri), (ci + grid_span - 1, ri)))
                    for extra in range(1, grid_span):
                        seen_spans.add(id(row.cells[min(ci + extra, len(row.cells)-1)]._tc))

                ci += grid_span

            # pad row to n_cols
            while len(row_data) < n_cols:
                row_data.append("")
            data.append(row_data[:n_cols])

        if not data:
            return None

        t = Table(data, colWidths=col_widths, repeatRows=1,
                  hAlign="LEFT", splitByRow=True)
        t.setStyle(TableStyle(style_cmds))
        return t

    # ── build story ────────────────────────────────────────────────────────

    from docx.text.paragraph import Paragraph as DocxPara
    from docx.table import Table as DocxTable

    block_items = list(_iter_block_items(doc))
    story       = []

    for i, item in enumerate(block_items):
        if isinstance(item, DocxPara):
            el = process_paragraph(item)
        else:
            el = process_table(item)
            if el:
                story.append(Spacer(1, 6))
        if el is not None:
            story.append(el)
        if el and isinstance(item, DocxTable):
            story.append(Spacer(1, 8))
        if i % 10 == 0:
            log(f"Processing... ({i+1}/{len(block_items)})")

    if not story:
        story.append(Paragraph("(empty document)", normal))

    pdf.build(story)


def _iter_block_items(doc):
    """Yield paragraphs and tables in document order."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxPara
    from docx.table import Table as DocxTable

    parent = doc.element.body
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxPara(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


# ═══════════════════════════════════════════════════════════
#  GUI
# ═══════════════════════════════════════════════════════════

MODES = [
    ("CSV / XLSX  ->  Image",  "sheet2img"),
    ("PDF  ->  Images",        "pdf2img"),
    ("Images  ->  PDF",        "img2pdf"),
    ("DOCX  ->  PDF",          "docx2pdf"),
]

IMG_FMTS = ["PNG", "JPG", "JPEG"]

BG     = "#F8F9FB"
DARK   = "#2E4057"
GRAY   = "#6B7280"
BORDER = "#E5E7EB"
GREEN  = "#16A34A"
RED    = "#DC2626"


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("File Converter")
        self.resizable(False, False)
        self.configure(bg=BG)

        self.mode      = tk.StringVar(value="sheet2img")
        self.img_fmt   = tk.StringVar(value="PNG")
        self.files     = []
        self.status_sv = tk.StringVar(value="")

        self._build()
        self._center(560, 560)

    def _center(self, w, h):
        self.update_idletasks()
        x = (self.winfo_screenwidth()  - w) // 2
        y = (self.winfo_screenheight() - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

    def _build(self):
        tk.Label(self, text="File Converter", font=("Helvetica", 18, "bold"),
                 bg=BG, fg="#1A1A2E").pack(pady=(26, 4))
        tk.Label(self, text="Choose a conversion, pick your files, and go",
                 font=("Helvetica", 10), bg=BG, fg=GRAY).pack(pady=(0, 18))

        tk.Label(self, text="Conversion type", font=("Helvetica", 10, "bold"),
                 bg=BG, fg="#374151", anchor="w").pack(fill="x", padx=28)

        grid = tk.Frame(self, bg=BG)
        grid.pack(fill="x", padx=28, pady=(8, 18))
        self.mode_btns = {}
        for i, (label, key) in enumerate(MODES):
            btn = tk.Button(grid, text=label, font=("Helvetica", 9),
                            relief="flat", cursor="hand2",
                            padx=10, pady=8, wraplength=110,
                            command=lambda k=key: self._set_mode(k))
            btn.grid(row=0, column=i, padx=(0, 8), sticky="ew")
            grid.columnconfigure(i, weight=1)
            self.mode_btns[key] = btn

        self.fmt_frame = tk.Frame(self, bg=BG)
        self.fmt_frame.pack(fill="x", padx=28, pady=(0, 14))
        tk.Label(self.fmt_frame, text="Output format",
                 font=("Helvetica", 10, "bold"),
                 bg=BG, fg="#374151", anchor="w").pack(anchor="w", pady=(0, 6))
        btn_row = tk.Frame(self.fmt_frame, bg=BG)
        btn_row.pack(anchor="w")
        self.fmt_btns = {}
        for fmt in IMG_FMTS:
            b = tk.Button(btn_row, text=fmt, font=("Helvetica", 9),
                          relief="flat", cursor="hand2", padx=16, pady=6,
                          command=lambda f=fmt: self._set_fmt(f))
            b.pack(side="left", padx=(0, 8))
            self.fmt_btns[fmt] = b
        self._set_fmt("PNG")

        tk.Frame(self, height=1, bg=BORDER).pack(fill="x", padx=28, pady=(0, 14))
        hdr = tk.Frame(self, bg=BG)
        hdr.pack(fill="x", padx=28)
        self.file_lbl = tk.Label(hdr, text="Selected files",
                                 font=("Helvetica", 10, "bold"),
                                 bg=BG, fg="#374151")
        self.file_lbl.pack(side="left")
        tk.Button(hdr, text="Clear", font=("Helvetica", 9),
                  relief="flat", bg=BORDER, fg=GRAY, padx=8, pady=2,
                  cursor="hand2", command=self._clear_files).pack(side="right")
        tk.Button(hdr, text="+ Add files", font=("Helvetica", 9),
                  relief="flat", bg=DARK, fg="#FFFFFF", padx=10, pady=2,
                  cursor="hand2", command=self._browse).pack(side="right", padx=(0, 8))

        list_frame = tk.Frame(self, bg=BG, bd=0)
        list_frame.pack(fill="both", padx=28, pady=(8, 0), expand=True)
        sb = tk.Scrollbar(list_frame)
        sb.pack(side="right", fill="y")
        self.listbox = tk.Listbox(list_frame, yscrollcommand=sb.set,
                                  font=("Helvetica", 9), relief="flat",
                                  bg="#FFFFFF", fg="#374151",
                                  selectbackground="#E0E7FF",
                                  highlightthickness=1,
                                  highlightbackground=BORDER,
                                  height=7)
        self.listbox.pack(side="left", fill="both", expand=True)
        sb.config(command=self.listbox.yview)

        tk.Frame(self, height=1, bg=BORDER).pack(fill="x", padx=28, pady=(14, 0))
        self.convert_btn = tk.Button(self, text="Convert",
                                     font=("Helvetica", 11, "bold"),
                                     relief="flat", bg=DARK, fg="#FFFFFF",
                                     activebackground="#1e2d3d",
                                     activeforeground="#FFFFFF",
                                     cursor="hand2", pady=11,
                                     command=self._start)
        self.convert_btn.pack(fill="x", padx=28, pady=(14, 0))

        self.status_lbl = tk.Label(self, textvariable=self.status_sv,
                                   font=("Helvetica", 9), bg=BG, fg=GRAY,
                                   wraplength=500, justify="center")
        self.status_lbl.pack(pady=(10, 0))

        style = ttk.Style()
        style.theme_use("default")
        style.configure("C.Horizontal.TProgressbar",
                        troughcolor=BORDER, background=DARK,
                        thickness=3, borderwidth=0)
        self.progress = ttk.Progressbar(self, mode="indeterminate",
                                        style="C.Horizontal.TProgressbar")
        self.progress.pack(fill="x", padx=28, pady=(6, 18))

        # initialise mode AFTER all widgets are built
        self._set_mode("sheet2img")

    def _set_mode(self, key):
        self.mode.set(key)
        self.files.clear()
        self.listbox.delete(0, tk.END)
        self.status_sv.set("")
        for k, b in self.mode_btns.items():
            b.config(bg=DARK if k == key else BORDER,
                     fg="#FFFFFF" if k == key else "#374151")
        needs_fmt = key in ("sheet2img", "pdf2img")
        if needs_fmt:
            self.fmt_frame.pack(fill="x", padx=28, pady=(0, 14),
                                before=self.convert_btn)
        else:
            self.fmt_frame.pack_forget()

    def _set_fmt(self, fmt):
        self.img_fmt.set(fmt)
        for f, b in self.fmt_btns.items():
            b.config(bg=DARK if f == fmt else BORDER,
                     fg="#FFFFFF" if f == fmt else "#374151")

    def _browse(self):
        mode = self.mode.get()
        if mode == "sheet2img":
            types = [("Spreadsheets", "*.csv *.xlsx *.xls")]
            multi = False
        elif mode == "pdf2img":
            types = [("PDF files", "*.pdf")]
            multi = False
        elif mode == "img2pdf":
            types = [("Images", "*.png *.jpg *.jpeg *.bmp *.tiff *.webp")]
            multi = True
        else:
            types = [("Word documents", "*.docx")]
            multi = False

        if multi:
            paths = filedialog.askopenfilenames(title="Choose files", filetypes=types)
            for p in paths:
                if p not in self.files:
                    self.files.append(p)
                    self.listbox.insert(tk.END, Path(p).name)
        else:
            path = filedialog.askopenfilename(title="Choose a file", filetypes=types)
            if path:
                self.files = [path]
                self.listbox.delete(0, tk.END)
                self.listbox.insert(tk.END, Path(path).name)

    def _clear_files(self):
        self.files.clear()
        self.listbox.delete(0, tk.END)
        self.status_sv.set("")

    def _log(self, msg):
        self.after(0, lambda: self.status_sv.set(msg))

    def _start(self):
        if not self.files:
            messagebox.showwarning("No files", "Please add at least one file first.")
            return
        self.convert_btn.config(state="disabled")
        self.progress.start(10)
        self.status_sv.set("Working...")
        self.status_lbl.config(fg=GRAY)
        threading.Thread(target=self._run, daemon=True).start()

    def _run(self):
        mode = self.mode.get()
        fmt  = self.img_fmt.get().lower()
        try:
            if mode == "sheet2img":
                for f in self.files:
                    spreadsheet_to_image(f, fmt, self._log)
            elif mode == "pdf2img":
                for f in self.files:
                    pdf_to_images(f, fmt, self._log)
            elif mode == "img2pdf":
                images_to_pdf(self.files, self._log)
            elif mode == "docx2pdf":
                for f in self.files:
                    docx_to_pdf(f, self._log)
            self.after(0, self._done, True, None)
        except Exception as e:
            self.after(0, self._done, False, str(e))

    def _done(self, ok, err):
        self.progress.stop()
        self.convert_btn.config(state="normal")
        if ok:
            self.status_sv.set("Done! Files saved in the same folder as input.")
            self.status_lbl.config(fg=GREEN)
        else:
            self.status_sv.set(f"Error: {err}")
            self.status_lbl.config(fg=RED)


if __name__ == "__main__":
    app = App()
    app.mainloop()
